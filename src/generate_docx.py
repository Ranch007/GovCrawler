"""
从 allKeyword.xlsx 批量生成 .docx / .pdf 政策文件。

独立运行:  python -m src.generate_docx [--fmt docx|pdf|both]
前置条件:  需先执行 govcrawler.py 完成数据采集 → output/allKeyword.xlsx

依赖 xlsx 列（与 src.json_handler.EXPORT_FIELDS 对应）:
    序号 | pcode | title | pubtimeStr | url | childtype | puborg | post | 匹配关键词
"""

import os
import re
import sys

from openpyxl import load_workbook

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import OUTPUT_DIR, ALLKEYWORD_XLSX, DOCX_DIR, PDF_DIR
from src.json_handler import EXPORT_FIELDS

# ── 本地配置 ──
TITLE_MAX_LEN = 60
# docx/pdf 需要的 xlsx 列（EXPORT_FIELDS 去掉 url）
_FIELDS_NEEDED = [f for f in EXPORT_FIELDS if f != "url"]
# 元数据标签（docx/pdf 共用）
_META_LABELS = ["主题分类", "发文机关", "标题", "发文字号", "发布日期", "关键词"]


# ═══════════════════════════════════════════════════════════════
# 工具函数
# ═══════════════════════════════════════════════════════════════

def _safe_str(val) -> str:
    """None → ''，其余 strip"""
    return str(val).strip() if val is not None else ""


def _split_paragraphs(text: str) -> list:
    """自然段落优先切分，兼容旧 sanitized 数据的句末标点回退"""
    if not text:
        return []
    if "\n\n" in text:
        # 新数据：保留 _extract_content 产出的自然段落
        paras = [p.strip() for p in text.split("\n\n") if p.strip()]
        return [ln for p in paras for ln in p.split("\n") if ln.strip()]
    # 旧数据：sanitized 无换行，按句末标点切段
    text = re.sub(r"([。！？])", r"\1\n", text)
    return [ln for ln in text.split("\n") if ln]


def _safe_filename(title: str, seq: str, fallback_idx: int, ext: str) -> str:
    """生成安全文件名：序号_标题.ext"""
    seq_str = seq.zfill(2) if seq else str(fallback_idx).zfill(2)
    if title:
        safe = re.sub(r'[\\/:*?"<>|]', "", title)
        safe = safe[:TITLE_MAX_LEN] if len(safe) > TITLE_MAX_LEN else safe
    else:
        safe = "无标题"
    return f"{seq_str}_{safe}.{ext}"


# ═══════════════════════════════════════════════════════════════
# DOCX 生成
# ═══════════════════════════════════════════════════════════════

def _generate_docx(items: list, output_dir: str) -> int:
    """items: [(seq, pcode, title, pubtime, topic, org, post, kw), ...]"""
    from docx import Document
    from docx.shared import Pt

    os.makedirs(output_dir, exist_ok=True)
    count = 0
    for seq, pcode, title, pubtime, topic, org, post, kw in items:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "仿宋"
        style.font.size = Pt(12)

        values = [topic, org, title, pcode, pubtime, kw.replace(", ", "、") if kw else ""]
        for i, (label, value) in enumerate(zip(_META_LABELS, values)):
            p = doc.add_paragraph()
            run = p.add_run(f"{label}：{value}")
            if i == 2:
                run.bold = True

        doc.add_paragraph()
        if post:
            for line in _split_paragraphs(post):
                doc.add_paragraph(line)
        filename = _safe_filename(title, seq, count + 1, "docx")
        filepath = os.path.join(output_dir, filename)
        try:
            doc.save(filepath)
            count += 1
            print(f"  [docx {count}] {filename}")
        except PermissionError:
            print(f"  [跳过] {filename} — 文件被占用")

    print(f"\n  [DOCX] {count} 篇 -> {output_dir}/")
    return count


# ═══════════════════════════════════════════════════════════════
# PDF 生成
# ═══════════════════════════════════════════════════════════════

def _find_chinese_font() -> str | None:
    """按优先级查找系统中可用的中文字体文件"""
    candidates = [
        # Windows
        r"C:\Windows\Fonts\simfang.ttf",     # 仿宋（与 docx 一致）
        r"C:\Windows\Fonts\simsun.ttc",      # 宋体
        r"C:\Windows\Fonts\msyh.ttc",        # 微软雅黑
        r"C:\Windows\Fonts\simhei.ttf",      # 黑体
        # macOS
        "/System/Library/Fonts/STFangsong.ttf",
        "/System/Library/Fonts/PingFang.ttc",
        # Linux
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    ]
    for path in candidates:
        if os.path.isfile(path):
            return path
    return None


# ── PDF 字体缓存（跨 FPDF 实例复用，避免每篇重复解析 ~7MB TTF） ──
_FONT_CACHE: dict[str, tuple] = {}  # font_path → (fonts, font_files, diffs)


def _generate_pdf(items: list, output_dir: str) -> int:
    """items: [(seq, pcode, title, pubtime, topic, org, post, kw), ...]"""
    import copy
    import warnings
    from fpdf import FPDF

    font_path = _find_chinese_font()
    if not font_path:
        print("\n  [PDF] 未找到中文字体文件，跳过 PDF 生成")
        print("  提示：Windows 用户请确认 C:\\Windows\\Fonts\\ 下有 simfang.ttf 或 simsun.ttc")
        return 0

    os.makedirs(output_dir, exist_ok=True)

    # 字体数据加载一次，后续实例 deepcopy 播种
    if font_path not in _FONT_CACHE:
        dummy = FPDF()
        dummy.add_font(family="CJK", fname=font_path, uni=True)
        _FONT_CACHE[font_path] = (
            copy.deepcopy(dummy.fonts),
            copy.deepcopy(dummy.font_files),
            copy.deepcopy(dummy.diffs),
        )

    count = 0
    for seq, pcode, title, pubtime, topic, org, post, kw in items:
        pdf = FPDF()
        # 从缓存播种字体，避免每个 PDF 重新解析 TTF
        pdf.fonts = copy.deepcopy(_FONT_CACHE[font_path][0])
        pdf.font_files = copy.deepcopy(_FONT_CACHE[font_path][1])
        pdf.diffs = copy.deepcopy(_FONT_CACHE[font_path][2])
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.add_page()

        values = [topic, org, title, pcode, pubtime, kw.replace(", ", "、") if kw else ""]
        with warnings.catch_warnings():
            warnings.filterwarnings("ignore", message=r"cmap value too big")
            for i, (label, value) in enumerate(zip(_META_LABELS, values)):
                pdf.set_font("CJK", "", 14 if i == 2 else 12)
                pdf.multi_cell(0, 8, f"{label}：{value}")

            pdf.ln(4)

            pdf.set_font("CJK", "", 12)
            if post:
                for line in _split_paragraphs(post):
                    pdf.multi_cell(0, 7, line)

            filename = _safe_filename(title, seq, count + 1, "pdf")
            filepath = os.path.join(output_dir, filename)
            try:
                pdf.output(filepath)
                count += 1
                print(f"  [pdf {count}] {filename}")
            except PermissionError:
                print(f"  [跳过] {filename} — 文件被占用")

    print(f"\n  [PDF] {count} 篇 -> {output_dir}/")
    return count


# ═══════════════════════════════════════════════════════════════
# 主流程
# ═══════════════════════════════════════════════════════════════

def main(fmt: str = "docx"):
    """
    读取 xlsx，按指定格式生成文件。

    Args:
        fmt: "docx" | "pdf" | "both"
    """
    xlsx_path = os.path.join(OUTPUT_DIR, ALLKEYWORD_XLSX)
    if not os.path.exists(xlsx_path):
        print(f"[FAIL] xlsx 文件不存在: {xlsx_path}")
        print(f"  请先运行 govcrawler.py 完成数据采集（需要 Phase 1–4）")
        sys.exit(1)

    wb = load_workbook(xlsx_path, data_only=True)
    if "政策数据" not in wb.sheetnames:
        print(f"[FAIL] xlsx 中未找到工作表 '政策数据'，可用工作表: {wb.sheetnames}")
        sys.exit(1)
    ws = wb["政策数据"]

    # ── 按表头名定位列索引（字段与 EXPORT_FIELDS 同步） ──
    header_row = [cell.value for cell in next(ws.iter_rows(min_row=1, max_row=1))]
    try:
        idx = {name: header_row.index(name) for name in _FIELDS_NEEDED}
    except ValueError as e:
        print(f"[FAIL] xlsx 表头缺少必要字段: {e}")
        sys.exit(1)

    # ── 构建数据列表（docx/pdf 共用） ──
    items = []
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, values_only=True):
        seq     = _safe_str(row[idx["序号"]])
        pcode   = _safe_str(row[idx["pcode"]])
        title   = _safe_str(row[idx["title"]])
        pubtime = _safe_str(row[idx["pubtimeStr"]])
        topic   = _safe_str(row[idx["childtype"]])
        org     = _safe_str(row[idx["puborg"]])
        post    = _safe_str(row[idx["post"]])
        kw      = _safe_str(row[idx["匹配关键词"]])
        if not title and not post:
            continue
        items.append((seq, pcode, title, pubtime, topic, org, post, kw))

    # ── 按格式生成 ──
    docx_count = pdf_count = 0
    if fmt in ("docx", "both"):
        docx_count = _generate_docx(items, os.path.join(OUTPUT_DIR, DOCX_DIR))
    if fmt in ("pdf", "both"):
        pdf_count = _generate_pdf(items, os.path.join(OUTPUT_DIR, PDF_DIR))

    parts = []
    if docx_count:
        parts.append(f"docx {docx_count} 篇 -> {OUTPUT_DIR}/{DOCX_DIR}/")
    if pdf_count:
        parts.append(f"pdf {pdf_count} 篇 -> {OUTPUT_DIR}/{PDF_DIR}/")
    print(f"完成！{'，'.join(parts)}")


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="xlsx 转 docx/pdf")
    parser.add_argument("--fmt", choices=["docx", "pdf", "both"], default="docx",
                        help="输出格式（默认 docx）")
    args = parser.parse_args()
    main(args.fmt)
